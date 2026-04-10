# -*- coding: utf-8 -*-
"""
《灵木求生》游戏策划案 - Word文档生成器
参考网易、腾讯、米哈游标准策划案格式
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            element = OxmlElement(f'w:{edge}')
            element.set(qn('w:val'), kwargs[edge].get('val', 'single'))
            element.set(qn('w:sz'), str(kwargs[edge].get('sz', 4)))
            element.set(qn('w:color'), kwargs[edge].get('color', '000000'))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def set_run_font(run, font_name='微软雅黑', font_size=11, bold=False, color=None):
    """设置run的字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level, style_name='Heading'):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return heading

def add_para(doc, text, font_size=11, bold=False, indent=False):
    """添加段落"""
    para = doc.add_paragraph()
    if indent:
        para.paragraph_format.first_line_indent = Cm(0.74)
    para.paragraph_format.line_spacing = 1.5
    run = para.add_run(text)
    set_run_font(run, font_size=font_size, bold=bold)
    return para

def add_table_with_style(doc, rows, cols, data=None):
    """添加带样式的表格"""
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    if data:
        for i, row_data in enumerate(data):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                cell = row.cells[j]
                cell.text = str(cell_text)
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.font.size = Pt(10)
                        run.font.name = '微软雅黑'
    return table

def create_game_design_doc():
    """创建完整的游戏策划案文档"""
    doc = Document()
    
    # ==================== 封面 ====================
    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(120)
    run = title.add_run('《灵木求生》游戏设计文档')
    set_run_font(run, font_size=28, bold=True)
    
    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('——东方玄幻末世求生游戏策划案 v2.0')
    set_run_font(run, font_size=18)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 基本信息表格
    info_table = doc.add_table(rows=6, cols=2)
    info_table.style = 'Table Grid'
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    info_data = [
        ('文档版本', 'V2.0'),
        ('游戏名称', '灵木求生（原：求生之开局一个小树屋）'),
        ('游戏类型', '2D像素风+水墨风格融合类游戏'),
        ('核心玩法', '塔防 + 肉鸽(Roguelike) + 生存建造 + 东方修仙'),
        ('改编来源', '起点小说《求生之开局一个小树屋》（作者：为何我是神）'),
        ('目标平台', '微信小游戏 + Web双端'),
    ]
    
    for i, (key, value) in enumerate(info_data):
        row = info_table.rows[i]
        row.cells[0].text = key
        row.cells[1].text = value
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(11)
                    run.font.name = '微软雅黑'
                    if cell == row.cells[0]:
                        run.font.bold = True
    
    # 设置表格宽度
    for row in info_table.rows:
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(12)
    
    doc.add_page_break()
    
    # ==================== 目录 ====================
    add_heading(doc, '目录', 1)
    
    toc_items = [
        '1. 项目概述',
        '2. 游戏定位与差异化',
        '3. 核心玩法设计',
        '4. 游戏系统详细设计',
        '5. 数值框架设计',
        '6. 商业化设计',
        '7. UI/UX设计',
        '8. 美术风格规范',
        '9. 音效与音乐',
        '10. 技术架构',
        '11. 开发计划',
    ]
    
    for item in toc_items:
        para = doc.add_paragraph(item)
        para.paragraph_format.left_indent = Cm(0.5)
        para.paragraph_format.line_spacing = 1.8
        for run in para.runs:
            run.font.size = Pt(12)
            run.font.name = '微软雅黑'
    
    doc.add_page_break()
    
    # ==================== 正文 ====================
    
    # ========== 1. 项目概述 ==========
    add_heading(doc, '1. 项目概述', 1)
    
    add_heading(doc, '1.1 项目背景', 2)
    add_para(doc, '《灵木求生》是一款基于起点小说《求生之开局一个小树屋》（作者：为何我是神）改编的东方玄幻末世求生游戏。游戏融合了猛鬼宿舍（塔防）、肉鸽（Roguelike）、生存建造和东方修仙元素，打造独具特色的修仙题材生存游戏。', indent=True)
    
    add_heading(doc, '1.2 核心卖点', 2)
    
    # 卖点表格
    sell_table = add_table_with_style(doc, 5, 2, [
        ('卖点', '描述'),
        ('东方玄幻题材', '以修仙文化为核心，区别于西方哥特风格，建立差异化竞争优势'),
        ('水墨像素美术', '融合2D像素风格与中国水墨画元素，打造独特视觉体验'),
        ('健康付费生态', '不逼氪不逼肝，强调公平竞技与内容体验优先'),
        ('修仙文化深度', '融入渡劫、功德、道心、问道等修仙元素，提供沉浸式修仙体验'),
    ])
    
    doc.add_paragraph()
    
    add_heading(doc, '1.3 目标用户', 2)
    add_para(doc, '• 16-35岁游戏玩家', indent=True)
    add_para(doc, '• 修仙/玄幻题材爱好者', indent=True)
    add_para(doc, '• 塔防类游戏玩家', indent=True)
    add_para(doc, '• 休闲放置类游戏用户', indent=True)
    add_para(doc, '• 小游戏平台用户', indent=True)
    
    # ========== 2. 游戏定位与差异化 ==========
    add_heading(doc, '2. 游戏定位与差异化', 1)
    
    add_heading(doc, '2.1 市场定位', 2)
    add_para(doc, '《灵木求生》定位于修仙题材末世求生细分市场，目标成为该品类的头部产品。与传统塔防游戏（如猛鬼宿舍/躺平发育）形成差异化竞争，主打东方文化和修仙深度。', indent=True)
    
    add_heading(doc, '2.2 差异化对比', 2)
    
    diff_table = add_table_with_style(doc, 5, 4, [
        ('维度', '躺平发育', '灵木求生', '差异化优势'),
        ('题材', '西方哥特/恐怖', '东方玄幻/修仙', '文化认同感更强'),
        ('美术', '写实恐怖风格', '水墨像素融合', '视觉独特性强'),
        ('战斗', '武器塔防', '阵法符箓', '策略深度更高'),
        ('成长', '科技升级', '修仙渡劫', '文化内涵丰富'),
    ])
    
    doc.add_paragraph()
    
    add_heading(doc, '2.3 核心定位', 2)
    add_para(doc, '"一款融合东方修仙文化的末世求生游戏，让玩家在修仙问道中体验塔防与肉鸽的乐趣。"', indent=True)
    
    # ========== 3. 核心玩法设计 ==========
    add_heading(doc, '3. 核心玩法设计', 1)
    
    add_heading(doc, '3.1 游戏循环', 2)
    add_para(doc, '游戏采用「修仙计时循环」为核心驱动的玩法结构：', indent=True)
    add_para(doc, '• 子时（23:00-01:00）：妖潮来袭，抵御妖兽进攻', indent=True)
    add_para(doc, '• 丑时-亥时：白天准备阶段，升级灵木、布置阵法、炼制符箓', indent=True)
    add_para(doc, '• 卯时（05:00-07:00）：天亮休整，结算奖励，触发天命劫事件', indent=True)
    
    add_heading(doc, '3.2 核心玩法模块', 2)
    
    # 玩法表格
    play_table = add_table_with_style(doc, 5, 3, [
        ('玩法模块', '核心机制', '修仙元素'),
        ('灵木仙府', '树屋建造与升级', '渡劫/灵根/枝叶变异'),
        ('妖兽抵御', '塔防战斗', '五行阵法/符箓/妖兽抵御'),
        ('天命劫', 'Roguelike随机事件', '命格/道心/问道'),
        ('问道修仙', '角色成长线', '筑基/金丹/元婴'),
    ])
    
    doc.add_paragraph()
    
    # ========== 4. 游戏系统详细设计 ==========
    add_heading(doc, '4. 游戏系统详细设计', 1)
    
    # 4.1 灵木仙府系统
    add_heading(doc, '4.1 灵木仙府系统（树屋建造系统）', 2)
    
    add_heading(doc, '4.1.1 系统概述', 3)
    add_para(doc, '灵木仙府系统是游戏的核心建造模块，将传统"树屋"升级为具有修仙特色的"灵木"。玩家通过培养灵木、凝聚灵根、经历渡劫来提升实力。', indent=True)
    
    add_heading(doc, '4.1.2 核心机制', 3)
    add_para(doc, '【灵木培养】', bold=True)
    add_para(doc, '• 基础属性：树高、树冠范围、灵气容量', indent=True)
    add_para(doc, '• 生长阶段：幼苗期→成长期→成熟期→化形期→渡劫期', indent=True)
    add_para(doc, '• 生长条件：灵气、水源、阳光（昼夜影响）', indent=True)
    
    add_para(doc, '【灵根系统】', bold=True)
    add_para(doc, '• 灵根品质：金木水火土五行灵根，品质从低到高：C/B/A/S/SS', indent=True)
    add_para(doc, '• 灵根觉醒：通过修炼觉醒灵根，每次渡劫可选择强化', indent=True)
    add_para(doc, '• 灵根共鸣：相同灵根可产生共鸣效果，提升属性', indent=True)
    
    add_para(doc, '【渡劫系统】', bold=True)
    add_para(doc, '• 渡劫触发：每10波妖兽进攻后触发渡劫事件', indent=True)
    add_para(doc, '• 渡劫准备：消耗材料炼制渡劫丹，准备阵法护体', indent=True)
    add_para(doc, '• 渡劫结果：成功则提升境界，失败则降级但保留部分属性', indent=True)
    
    add_heading(doc, '4.1.3 枝叶变异机制', 3)
    add_para(doc, '• 变异类型：攻击型、防御型、辅助型、特殊型', indent=True)
    add_para(doc, '• 变异条件：灵气浓郁度、渡劫次数、特殊事件触发', indent=True)
    add_para(doc, '• 变异效果：提供额外属性加成和特殊技能', indent=True)
    
    add_heading(doc, '4.1.4 数值设计', 3)
    
    lingmu_table = add_table_with_style(doc, 6, 5, [
        ('境界', '树高(m)', '灵气上限', '渡劫成功率', '解锁功能'),
        ('炼气期', '5-10', '100', '80%', '基础建造'),
        ('筑基期', '10-20', '500', '70%', '灵根觉醒'),
        ('金丹期', '20-35', '2000', '60%', '渡劫阵法'),
        ('元婴期', '35-50', '8000', '50%', '化形技能'),
        ('化神期', '50+', '30000', '40%', '完全化形'),
    ])
    
    doc.add_paragraph()
    
    # 4.2 妖兽抵御系统
    add_heading(doc, '4.2 妖兽抵御系统（防御战斗系统）', 2)
    
    add_heading(doc, '4.2.1 系统概述', 3)
    add_para(doc, '妖兽抵御系统是游戏的战斗核心，将传统"武器塔"替换为具有东方特色的"阵法符箓"系统。玩家通过布置阵法、炼制符箓来抵御妖兽进攻。', indent=True)
    
    add_heading(doc, '4.2.2 阵法系统', 3)
    add_para(doc, '【五行阵法】', bold=True)
    
    wuxing_table = add_table_with_style(doc, 6, 4, [
        ('阵法', '相克', '被克', '特效'),
        ('金剑阵', '木', '火', '高额单体伤害'),
        ('木缠绕阵', '土', '金', '控制敌人移动'),
        ('水冰阵', '火', '土', '范围减速'),
        ('火烈阵', '金', '水', '持续灼烧'),
        ('土厚阵', '水', '木', '高额防御'),
    ])
    
    doc.add_paragraph()
    add_para(doc, '【阵法升级】', bold=True)
    add_para(doc, '• 基础升级：消耗灵石提升阵法等级，增加伤害/效果', indent=True)
    add_para(doc, '• 阵法共鸣：相邻同属性阵法产生共鸣，额外加成', indent=True)
    add_para(doc, '• 阵法融合：高级玩家可将两种阵法融合，产生复合效果', indent=True)
    
    add_heading(doc, '4.2.3 符箓系统', 3)
    add_para(doc, '• 攻击符箓：雷符（火雷/水雷/金雷）、剑符', indent=True)
    add_para(doc, '• 防御符箓：护体符、金钟符、护盾符', indent=True)
    add_para(doc, '• 控制符箓：定身符、困兽符、迷雾符', indent=True)
    add_para(doc, '• 辅助符箓：治疗符、灵气符、聚灵符', indent=True)
    
    add_heading(doc, '4.2.4 妖兽设计', 3)
    
    monster_table = add_table_with_style(doc, 8, 5, [
        ('妖兽名称', '等级', '属性', '弱点', '掉落'),
        ('游魂', '1-5', '阴', '阳/雷', '灵气'),
        ('野狼妖', '5-10', '风', '土/火', '妖丹'),
        ('狐妖', '10-20', '魅', '金/火', '媚珠'),
        ('蛇妖', '15-25', '毒', '火/雷', '毒囊'),
        ('蛟龙', '25-35', '水', '土/雷', '龙鳞'),
        ('饕餮', '35-45', '贪', '无', '饕餮血'),
        ('妖帝', '45+', '混合', '组合', '大量资源'),
    ])
    
    doc.add_paragraph()
    
    # 4.3 天命劫系统
    add_heading(doc, '4.3 天命劫系统（肉鸽随机事件系统）', 2)
    
    add_heading(doc, '4.3.1 系统概述', 3)
    add_para(doc, '天命劫系统是游戏的Roguelike核心，将传统"随机事件"升级为具有修仙命格的"天道轮回"机制。玩家通过问道修仙、考验道心来获得额外奖励和成长。', indent=True)
    
    add_heading(doc, '4.3.2 命格系统', 3)
    add_para(doc, '【命格类型】', bold=True)
    
    mingge_table = add_table_with_style(doc, 6, 3, [
        ('命格类型', '获取方式', '效果'),
        ('天道眷顾', '天命事件', '渡劫成功率+10%'),
        ('五行亲和', '灵根觉醒', '对应属性伤害+15%'),
        ('大气运者', '稀有事件', '资源掉落+20%'),
        ('多灾多难', '负面事件', '属性-10%但概率触发奇遇'),
        ('因果缠身', '击杀妖兽', '功德值影响命格变化'),
    ])
    
    doc.add_paragraph()
    
    add_heading(doc, '4.3.3 道心考验', 3)
    add_para(doc, '• 道心品质：坚固→坚定→坚如磐石→不动如山→天人合一', indent=True)
    add_para(doc, '• 考验类型：心魔试炼、功德抉择、问道求索、渡劫考验', indent=True)
    add_para(doc, '• 考验奖励：道心越高，获得的属性加成和特殊能力越强', indent=True)
    
    add_heading(doc, '4.3.4 问道修仙', 3)
    add_para(doc, '【问道方式】', bold=True)
    add_para(doc, '• 问道于天：通过天降异象获取天机提示', indent=True)
    add_para(doc, '• 问道于地：通过探索秘境获取机缘', indent=True)
    add_para(doc, '• 问道于师：通过仙盟拜师获取传承', indent=True)
    add_para(doc, '• 问道于心：通过内观自省提升道心', indent=True)
    
    # 4.4 灵兽仙宠系统
    add_heading(doc, '4.4 灵兽仙宠系统', 2)
    
    add_heading(doc, '4.4.1 系统概述', 3)
    add_para(doc, '灵兽仙宠系统为玩家提供战斗伙伴，玩家可以捕捉野生灵兽，通过培养提升实力，最终契约仙宠。', indent=True)
    
    add_heading(doc, '4.4.2 灵兽品质', 3)
    
    spirit_table = add_table_with_style(doc, 6, 4, [
        ('品质', '原型', '技能数', '成长潜力'),
        ('凡品', '野兔/野鸡', '1', '低'),
        ('良品', '狼/狐狸', '2', '中'),
        ('上品', '白虎/玄蛇', '3', '高'),
        ('极品', '朱雀/青龙', '4', '极高'),
        ('神品', '麒麟/鲲鹏', '5', '极限'),
    ])
    
    doc.add_paragraph()
    add_para(doc, '【契约系统】', bold=True)
    add_para(doc, '• 灵兽蛋：野外概率获得，放入灵泉孵化', indent=True)
    add_para(doc, '• 灵兽培养：消耗灵气和灵草提升等级和技能', indent=True)
    add_para(doc, '• 灵兽进化：达成条件后进化为更高品质', indent=True)
    add_para(doc, '• 仙宠契约：最高阶灵兽可签订契约，获得专属技能', indent=True)
    
    # 4.5 仙盟系统
    add_heading(doc, '4.5 仙盟系统', 2)
    
    add_heading(doc, '4.5.1 仙盟创建', 3)
    add_para(doc, '• 创建条件：等级达到筑基期，消耗5000灵石', indent=True)
    add_para(doc, '• 仙盟名称：2-6个汉字，不可重复', indent=True)
    add_para(doc, '• 仙盟宣言：30字内的宗旨宣言', indent=True)
    
    add_heading(doc, '4.5.2 仙盟功能', 3)
    add_para(doc, '• 仙盟商店：使用仙盟贡献兑换稀有道具', indent=True)
    add_para(doc, '• 仙盟任务：每日任务获取贡献和灵石', indent=True)
    add_para(doc, '• 仙盟副本：组队挑战妖王，获取高级材料', indent=True)
    add_para(doc, '• 仙盟技能：集体增益技能，全盟生效', indent=True)
    add_para(doc, '• 仙盟排名：按贡献和实力排名，赛季奖励', indent=True)
    
    # 4.6 装备系统
    add_heading(doc, '4.6 装备系统', 2)
    
    add_heading(doc, '4.6.1 装备类型', 3)
    
    equip_table = add_table_with_style(doc, 6, 3, [
        ('装备部位', '名称', '属性加成'),
        ('武器', '法剑/法杖/拂尘', '攻击、暴击'),
        ('防具', '道袍/法衣/仙袍', '防御、生命'),
        ('饰品', '玉佩/戒指/手镯', '特殊属性'),
        ('法器', '葫芦/宝塔/铜镜', '技能加成'),
        ('坐骑', '仙鹤/白鹿/云鹏', '移动速度、灵气恢复'),
    ])
    
    doc.add_paragraph()
    
    add_heading(doc, '4.6.2 装备品质', 3)
    add_para(doc, '白色（普通）→ 绿色（良品）→ 蓝色（上品）→ 紫色（极品）→ 橙色（仙器）→ 红色（神器）', indent=True)
    
    add_heading(doc, '4.6.3 装备强化', 3)
    add_para(doc, '• 灵石强化：消耗灵石提升装备等级', indent=True)
    add_para(doc, '• 炼器升级：消耗材料进行品质突破', indent=True)
    add_para(doc, '• 套装效果：集齐套装激活额外属性', indent=True)
    
    # ========== 5. 数值框架设计 ==========
    add_heading(doc, '5. 数值框架设计', 1)
    
    add_heading(doc, '5.1 成长曲线', 2)
    add_para(doc, '游戏采用「对数成长+台阶式突破」的数值模型，确保玩家在前期快速成长获得正反馈，中后期通过境界突破获得成就感。', indent=True)
    
    add_heading(doc, '5.2 货币系统', 3)
    
    currency_table = add_table_with_style(doc, 5, 3, [
        ('货币类型', '获取方式', '主要用途'),
        ('灵石', '妖兽掉落、任务奖励', '建造升级、阵法升级'),
        ('灵晶', '充值、月卡', '商城购买、限时道具'),
        ('功德', '帮助他人、积累善行', '功德商店、好感度'),
        ('仙缘', '稀有事件、问道求索', '抽取仙宠、稀有功法'),
    ])
    
    doc.add_paragraph()
    
    add_heading(doc, '5.3 资源产出模型', 3)
    
    resource_table = add_table_with_style(doc, 4, 5, [
        ('阶段', '灵石产出', '灵气消耗', '升级周期', '难度曲线'),
        ('炼气期', '100/波', '50/波', '3-5分钟', '平滑'),
        ('筑基期', '300/波', '200/波', '5-10分钟', '略陡'),
        ('金丹期', '800/波', '600/波', '10-15分钟', '陡峭'),
    ])
    
    doc.add_paragraph()
    
    add_heading(doc, '5.4 数值平衡原则', 3)
    add_para(doc, '• 克制关系：五行相克伤害+50%，被克伤害-50%', indent=True)
    add_para(doc, '• 难度递进：每10波难度提升20%，渡劫期难度翻倍', indent=True)
    add_para(doc, '• 付费平衡：付费道具提供便利但不提供数值碾压', indent=True)
    add_para(doc, '• 生态健康：免费玩家可通过时间和技巧追平付费差距', indent=True)
    
    # ========== 6. 商业化设计 ==========
    add_heading(doc, '6. 商业化设计', 1)
    
    add_heading(doc, '6.1 付费理念', 2)
    add_para(doc, '"不逼氪不逼肝，让玩家为快乐付费，而非为焦虑买单"', indent=True)
    add_para(doc, '• 拒绝Pay-to-Win：核心数值不可购买', indent=True)
    add_para(doc, '• 拒绝疲劳值：可随时退出，不强制在线', indent=True)
    add_para(doc, '• 拒绝套路：明码标价，无隐藏概率', indent=True)
    
    add_heading(doc, '6.2 付费点设计', 3)
    
    pay_table = add_table_with_style(doc, 6, 4, [
        ('付费类型', '定价区间', '内容', '价值感知'),
        ('月卡', '¥30/月', '每日100灵晶+特权', '高性价比'),
        ('战令', '¥68/赛季', '专属皮肤+道具', '内容丰富'),
        ('直购', '¥6-648', '皮肤/道具/资源', '明码标价'),
        ('基金', '¥68', '投资返还+额外奖励', '长期价值'),
        ('礼包', '¥6-128', '限时优惠组合', '促销驱动'),
    ])
    
    doc.add_paragraph()
    
    add_heading(doc, '6.3 付费安全性', 3)
    add_para(doc, '• 未成年人保护：严格实名认证，限时消费', indent=True)
    add_para(doc, '• 消费冷静期：大额消费二次确认', indent=True)
    add_para(doc, '• 退款机制：合规渠道退款申请', indent=True)
    
    # ========== 7. UI/UX设计 ==========
    add_heading(doc, '7. UI/UX设计', 1)
    
    add_heading(doc, '7.1 设计原则', 2)
    add_para(doc, '• 一致性：保持修仙风格的视觉统一', indent=True)
    add_para(doc, '• 可读性：重要信息一目了然', indent=True)
    add_para(doc, '• 反馈性：操作即时反馈，降低认知负担', indent=True)
    add_para(doc, '• 可达性：单手操作适配', indent=True)
    
    add_heading(doc, '7.2 界面布局', 3)
    
    ui_table = add_table_with_style(doc, 5, 2, [
        ('界面', '布局'),
        ('主界面', '顶部资源栏+中央游戏区+底部快捷栏'),
        ('建造界面', '左侧分类栏+中央物品栏+右侧详情面板'),
        ('战斗界面', '顶部波次信息+中央战场+底部技能栏'),
        ('修仙界面', '角色立绘+属性面板+技能树'),
    ])
    
    doc.add_paragraph()
    
    add_heading(doc, '7.3 交互规范', 3)
    add_para(doc, '• 点击响应时间：<100ms', indent=True)
    add_para(doc, '• 手势支持：点击、长按、滑动、拖拽', indent=True)
    add_para(doc, '• 震动反馈：重要事件震动提醒', indent=True)
    
    # ========== 8. 美术风格规范 ==========
    add_heading(doc, '8. 美术风格规范', 1)
    
    add_heading(doc, '8.1 整体风格', 2)
    add_para(doc, '《灵木求生》采用「水墨像素融合」的独特美术风格，将传统水墨画的写意神韵与现代像素艺术的表现力相结合。', indent=True)
    
    add_heading(doc, '8.2 色彩规范', 3)
    
    color_table = add_table_with_style(doc, 5, 2, [
        ('色彩', '色值'),
        ('主色调（灵木绿）', '#4A7C59'),
        ('辅助色（仙气紫）', '#9B7ED9'),
        ('强调色（雷电金）', '#FFD700'),
        ('背景色（云雾白）', '#F5F5F5'),
    ])
    
    doc.add_paragraph()
    
    add_heading(doc, '8.3 视觉资产', 3)
    add_para(doc, '• 角色立绘：仙侠古风，2D手绘风格', indent=True)
    add_para(doc, '• 场景原画：水墨山水意境，层次分明', indent=True)
    add_para(doc, '• 技能特效：粒子+光效+水墨笔触', indent=True)
    add_para(doc, '• UI图标：统一像素风格，辨识度高', indent=True)
    
    # ========== 9. 音效与音乐 ==========
    add_heading(doc, '9. 音效与音乐', 1)
    
    add_heading(doc, '9.1 音效设计', 2)
    add_para(doc, '• 界面音效：点击、滑动、确认、取消', indent=True)
    add_para(doc, '• 游戏音效：建造、攻击、技能、击杀', indent=True)
    add_para(doc, '• 环境音效：风声、雨声、鸟鸣、妖兽叫声', indent=True)
    add_para(doc, '• 反馈音效：升级、渡劫成功、获得稀有物品', indent=True)
    
    add_heading(doc, '9.2 音乐设计', 2)
    add_para(doc, '• 主界面音乐：悠扬的古筝+笛子，宁静祥和', indent=True)
    add_para(doc, '• 战斗音乐：激昂的鼓点+琵琶，紧张刺激', indent=True)
    add_para(doc, '• 渡劫音乐：庄重的编钟+古琴，肃穆神圣', indent=True)
    add_para(doc, '• 胜利音乐：欢快的丝竹+打击乐，喜庆热闹', indent=True)
    
    # ========== 10. 技术架构 ==========
    add_heading(doc, '10. 技术架构', 1)
    
    add_heading(doc, '10.1 技术选型', 2)
    
    tech_table = add_table_with_style(doc, 6, 2, [
        ('模块', '技术方案'),
        ('游戏引擎', 'Phaser.js 3.x'),
        ('渲染方式', 'Canvas 2D/WebGL'),
        ('网络通信', 'WebSocket'),
        ('数据存储', 'LocalStorage + 云端服务器'),
        ('适配方案', '响应式布局，支持多分辨率'),
    ])
    
    doc.add_paragraph()
    
    add_heading(doc, '10.2 性能目标', 3)
    add_para(doc, '• 帧率：稳定60FPS', indent=True)
    add_para(doc, '• 首屏加载：<3秒', indent=True)
    add_para(doc, '• 包体大小：<5MB（小游戏）', indent=True)
    add_para(doc, '• 内存占用：<200MB', indent=True)
    
    # ========== 11. 开发计划 ==========
    add_heading(doc, '11. 开发计划', 1)
    
    add_heading(doc, '11.1 MVP阶段（炼气期，5周）', 2)
    
    mvp_table = add_table_with_style(doc, 7, 3, [
        ('周次', '内容', '产出'),
        ('第1周', '核心架构搭建', '游戏框架、可运行Demo'),
        ('第2周', '灵木仙府基础功能', '建造系统、基础属性'),
        ('第3周', '阵法符箓战斗系统', '阵法放置、符箓使用'),
        ('第4周', '妖兽AI与波次系统', '妖兽生成、攻击逻辑'),
        ('第5周', '基础经济与UI', '灵石系统、界面搭建'),
        ('验收', '可玩版本测试', '核心玩法验证'),
    ])
    
    doc.add_paragraph()
    
    add_heading(doc, '11.2 后续迭代规划', 2)
    add_para(doc, '• 第6-10周：筑基期功能（灵根、渡劫、符箓系统）', indent=True)
    add_para(doc, '• 第11-15周：金丹期功能（问道、命格、灵兽）', indent=True)
    add_para(doc, '• 第16-20周：仙盟系统与社交功能', indent=True)
    add_para(doc, '• 第21-25周：商业化与运营功能', indent=True)
    
    # 保存文档
    output_path = os.path.join(os.path.dirname(__file__), '灵木求生游戏策划案v2.0.docx')
    doc.save(output_path)
    print(f'文档已生成: {output_path}')
    return output_path

if __name__ == '__main__':
    create_game_design_doc()
